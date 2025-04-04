import os
import comtypes.client
from docx import Document
from docx.shared import Inches
from PIL import Image

# Define paths
base_dir = os.path.dirname(os.path.abspath(__file__))
files_dir = os.path.join(base_dir, "Files")
pptx_path = os.path.join(files_dir, "ppt-to-doc-input.pptx")
image_output_folder = os.path.join(files_dir, "images")
docx_output_path = os.path.join(files_dir, "ppt-to-doc-output.docx")

# Create image output folder if it doesn't exist
os.makedirs(image_output_folder, exist_ok=True)

# Start PowerPoint (Windows only)
ppt_app = comtypes.client.CreateObject("PowerPoint.Application")
ppt_app.Visible = 1

# Open and export presentation as JPGs
presentation = ppt_app.Presentations.Open(pptx_path, WithWindow=False)
presentation.SaveAs(image_output_folder, 17)  # 17 = ppSaveAsJPG
presentation.Close()
ppt_app.Quit()

print("Slides exported to images.")

# Create Word document
doc = Document()

# Get and sort slide images (Slide1.JPG, Slide2.JPG, ...)
image_files = sorted(
    [f for f in os.listdir(image_output_folder) if f.lower().endswith('.jpg')],
    key=lambda name: int(''.join(filter(str.isdigit, name)))
)

for img_file in image_files:
    image_path = os.path.join(image_output_folder, img_file)
    with Image.open(image_path) as img:
        doc.add_picture(image_path, width=Inches(6))  # Resize to fit page
    doc.add_paragraph()  # Optional space between images

doc.save(docx_output_path)
print(f"DOCX created at {docx_output_path}")
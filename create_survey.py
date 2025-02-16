import pandas as pd
from docx import Document

intro_text = """¡Bienvenido(a)!

En el marco de la elaboración del Plan Estratégico de Internacionalización, Intercul-
turalidad y Relacionamiento de nuestra universidad, es esencial contar con informa-
ción precisa y relevante desde cada uno de los vicerrectorados y direcciones. Reco-
nocendo que cada dependencia tiene características particulares, solicitamos su 
participación en este cuestionário, con la participación activa y colaboración de su 
equipo, hasta 13 de febrero. Necesitamos una única respuesta por vicerrectorado/dirección.

Los datos recopilados se utilizarán exclusivamente con fines de mejora de nuestros 
servicios y serán tratados con total confidencialidad, de acuerdo con las directrices 
de la Ley General de Protección de Datos (LGPD). La identidad de los participantes 
será protegida en todas las etapas del proceso. Para más información sobre la LGPD, 
acceda a: www.planalto.gov.br/ccivil 03/_ato2015-2018/2018/lei/113709.htm.

En caso de dudas, pedimos entrar en contactar con nuestra consultoria IHub —
ihub@ihubeducacional.com.br."""

# Variables
input_file = 'survey_input.xlsx'
output_file = 'output.docx'
doc = Document()

# Load the Excel file with a multi-line header (first row = column title, second row = flag)
df = pd.read_excel(input_file, header=[0,1])

# Get the total number of columns and rows
total_columns = len(df.columns)
total_rows = len(df.index)

# Print basic file information
print(f"The file has {total_rows} rows and {total_columns} columns)")
print(f"Processing ...")

# Add the document title and intro text
doc.add_paragraph(df.columns[0][0])
doc.add_paragraph(intro_text)

# Extract relevant columns from the DataFrame
col_order = df.iloc[:, 0].values
col_section = df.iloc[:, 1].values
col_question = df.iloc[:, 2].values
col_answers = df.iloc[:, 3].fillna('').values

# Initialize the first section
current_section = col_section[0]
doc.add_paragraph(current_section)

# Looping through each row.
for row_index in range (total_rows):
    
    # If the section changes, add a new section heading
    if current_section != col_section[row_index]:
        current_section = col_section[row_index]
        doc.add_paragraph(current_section)

    # Add the question number and text to the document
    doc.add_paragraph(f"{col_order[row_index]}) {col_question[row_index]}")

    # Add the answer choices (if available)
    doc.add_paragraph(f"{col_answers[row_index]}")

# Save the Word document
doc.save(output_file)

print(f"Process completed successfully.")
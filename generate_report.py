import os, re
import pandas as pd
from docx import Document
from docx.shared import Inches
from ai_integration import ask_ai, get_report_building_prompt, get_section_analyzer_prompt
from config import INPUT_XLSX, LLM_FEATURES_ON, CONTROL_SHEET_NAME, GROUP_BY_COL_INDEX, GENERAL_LABEL
from llm_prompts import LLM_OUTPUT_SIMPLE, LLM_OUTPUT_BY_GROUPS, LLM_OUTPUT_SECTION_ANALYSIS
from summarizer import chart_data_map

# Parses the LLM output in Markdown into DOCX paraghraphs
def add_markdown_paragraph(doc, text):
    """
    Adds a paragraph to the docx Document, parsing **bold** and _italic_ markdown.
    """
    p = doc.add_paragraph()

    # Regex to capture bold (**...**) and italic (*...*) markers
    tokens = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', text)

    for token in tokens:
        if not token:
            continue
        run = p.add_run()

        if token.startswith("**") and token.endswith("**"):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith("__") and token.endswith("__"):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith("_") and token.endswith("_"):
            run.text = token[1:-1]
            run.italic = True
        else:
            run.text = token

# This function inserts the chart for a given column into the document
def insert_chart_for_column(doc, col_letter, charts_dir="charts"):
    for fname in os.listdir(charts_dir):
        if f"_column{col_letter}.jpg" in fname:
            chart_path = os.path.join(charts_dir, fname)
            doc.add_picture(chart_path, width=Inches(5.5))
            return True
    return False

# This function generates a .DOCX report by leveraging LLM
def generate_diagnosis_report():
    # Load Excel to detect grouping
    df_full = pd.read_excel(INPUT_XLSX, sheet_name=CONTROL_SHEET_NAME, header=1)

    # Get outline once
    prompt = get_report_building_prompt()
    print(f"LLM PROMPT:\n{prompt}\n")

    # Gets the output from LLM
    if LLM_FEATURES_ON:
        outline = ask_ai(prompt)
        print(f"\n LLM Output:{outline}\n")
    else:
        print("LLM disabled, using offline output")
        outline = LLM_OUTPUT_BY_GROUPS

    # If we want to split by groups (like schools or cities), we do that here
    if GROUP_BY_COL_INDEX is None:
        # Single global report
        build_report(outline, GENERAL_LABEL, charts_dir="charts")
    
    else:
        # One report per group
        group_col_name = df_full.columns[GROUP_BY_COL_INDEX]
        groups_df = df_full[df_full[group_col_name].notna() & (df_full[group_col_name] != "")]
        unique_groups = groups_df[group_col_name].unique().tolist()

        for g in unique_groups:
            group_label = str(g).strip() or "Unknown"
            charts_dir = "charts"  # still flat, but filenames include group name prefix
            build_report(outline, group_label, charts_dir)

def build_report(outline, group_label, charts_dir="charts"):
    doc = Document()
    doc.add_heading(f"Relatório — {group_label}", level=0)

    current_section = None
    current_cols = []

    for line in outline.split("\n"):
        if not line.strip():
            continue

        clean_line = line.strip()

        # --- Section headers
        if clean_line.startswith("## "):
            # Before starting a new section, flush previous one
            if current_section and current_cols:
                section_data = {c: chart_data_map.get(c, {}) for c in current_cols}
            
            if LLM_FEATURES_ON:
                analysis_text = ask_ai(get_section_analyzer_prompt(current_section, section_data))
            else:
                analysis_text = "(Análise automática desativada nesta execução)"
            doc.add_paragraph(analysis_text)

            # Start new section
            current_section = clean_line.replace("##", "").strip()
            current_cols = []
            doc.add_heading(current_section, level=2)
            continue

        elif clean_line.startswith("# "):
            doc.add_heading(clean_line[2:], level=1)
            continue

        elif clean_line.startswith("### "):
            doc.add_heading(clean_line[4:], level=3)
            continue

        else:
            # Parses the LLM Markdown output and inserts into the document
            add_markdown_paragraph(doc, clean_line)

            # Matches the placeholder for charts - so we can include it into the document
            col_match = re.match(r"[-•\s]*\*{0,2}([A-Z]{1,3})\*{0,2}\s*[:—\-–]", clean_line)
            if col_match:
                col_letter = col_match.group(1)
                if insert_chart_for_column(doc, col_letter, charts_dir):
                    current_cols.append(col_letter)

    # --- Flush last section
    if current_section and current_cols:
        section_data = {c: chart_data_map.get(c, {}) for c in current_cols}
        
        if LLM_FEATURES_ON:
            analysis_text = ask_ai(get_section_analyzer_prompt(current_section, section_data))

        else:
            analysis_text = LLM_OUTPUT_SECTION_ANALYSIS
            
        doc.add_paragraph(analysis_text)

    # build the report path
    base_name = os.path.splitext(os.path.basename(INPUT_XLSX))[0]
    safe_group = re.sub(r"[^\w\-]+", "_", group_label)
    report_path = os.path.join(os.path.dirname(INPUT_XLSX), f"{base_name}_{safe_group}_report.docx")

    # remove existing file if present
    if os.path.exists(report_path):
        os.remove(report_path)

    # Saves the report
    doc.save(report_path)
    print(f"Report saved to {report_path}")

if __name__ == "__main__":
    generate_diagnosis_report()